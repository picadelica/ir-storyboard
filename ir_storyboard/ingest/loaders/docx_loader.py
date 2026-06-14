"""Loader for .docx files exported from ChatGPT Deep Research / Claude / Gemini / Perplexity."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from ..ir import LLMReportIR, RawCitation, RawSection

# ── citation format regexes ──────────────────────────────────────────────────

_RE_BRACKET_N = re.compile(r"\[(\d+)\]")            # [1], [2], ...
_RE_PAREN_N = re.compile(r"\(Source\s+(\d+)\)", re.I)  # (Source 1)
_RE_SUPERSCRIPT = re.compile(r"[¹²³⁴-⁹]|[¹]")  # Unicode superscripts ¹²³...

# Superscript digit mapping
_SUPER_MAP = str.maketrans("¹²³⁴⁵⁶⁷⁸⁹⁰", "1234567890")

# ── agent detection ──────────────────────────────────────────────────────────

_AGENT_HINTS = {
    "chatgpt-deep-research": ["deep research", "openai", "chatgpt"],
    "claude": ["claude", "anthropic"],
    "perplexity": ["perplexity"],
    "gemini": ["gemini", "google"],
}


def _detect_agent(full_text: str) -> str | None:
    low = full_text[:2000].lower()
    for agent, hints in _AGENT_HINTS.items():
        if any(h in low for h in hints):
            return agent
    return "unknown"


def _detect_cite_format(full_text: str) -> str:
    if _RE_BRACKET_N.search(full_text):
        return "bracket_n"
    if _RE_PAREN_N.search(full_text):
        return "paren_source_n"
    if re.search(r"[¹²³⁴⁵⁶⁷⁸⁹]", full_text):
        return "superscript"
    return "unknown"


# ── URL extraction from citation list paragraphs ─────────────────────────────

_RE_URL = re.compile(r"https?://[^\s\]\)\>\"']+")

_RE_CITE_LINE_BRACKET = re.compile(
    r"^\[(\d+)\]\s*(.*?)(?:\s+[-–—]\s*|\s+)?(https?://\S+)?$"
)

# Hyperlink XML namespace
_HYPERLINK_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _para_text(para) -> str:
    """Extract full text from a paragraph, including runs."""
    return para.text or ""


def _para_hyperlinks(para) -> list[str]:
    """Extract all hyperlink URLs from a paragraph's XML."""
    urls = []
    for rel in para.part.rels.values():
        if "hyperlink" in rel.reltype:
            urls.append(rel._target)
    return urls


def _get_para_url(para) -> str:
    """Try to get the first hyperlink URL from a paragraph."""
    # Walk the XML tree for hyperlink relationships
    for child in para._element.iter():
        if child.tag.endswith("}hyperlink"):
            r_id = child.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
            )
            if r_id and r_id in para.part.rels:
                return para.part.rels[r_id]._target
    # Fallback: regex on text
    m = _RE_URL.search(para.text or "")
    return m.group(0) if m else ""


_REL_ID_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"


def _extract_inline_citations(para, out: dict[int, RawCitation]) -> None:
    """Harvest inline ``[N]`` citation markers that are themselves hyperlinks.

    ChatGPT Deep Research (and similar) exports attach the source URL directly
    to the ``[N]`` marker run in the body — there is no trailing Sources block.
    Each such hyperlink whose text matches ``[N]`` maps cite_id N → its URL.
    First URL seen for a given id wins (markers repeat across the body).
    """
    for child in para._element.iter():
        if not child.tag.endswith("}hyperlink"):
            continue
        r_id = child.get(_REL_ID_ATTR)
        if not r_id or r_id not in para.part.rels:
            continue
        link_text = "".join(
            t.text or "" for t in child.iter() if t.tag.endswith("}t")
        )
        m = _RE_BRACKET_N.search(link_text)
        if not m:
            continue
        cite_id = int(m.group(1))
        if cite_id in out and out[cite_id].url:
            continue
        url = (para.part.rels[r_id]._target or "").rstrip(".,)>")
        if not url:
            continue
        out[cite_id] = RawCitation(
            cite_id=cite_id,
            raw_marker=f"[{cite_id}]",
            url=url,
            title=out[cite_id].title if cite_id in out else "",
        )


# ── section heading detection ────────────────────────────────────────────────

_OPEN_QUESTIONS_HINTS = [
    "open questions", "open question", "вопросы для интервью",
    "вопросы к интервью", "открытые вопросы", "interview questions",
]

_SOURCES_HINTS = [
    "источники", "sources", "references", "bibliography", "ссылки",
]

_SKIP_SECTIONS = [
    "выводы", "conclusions", "заключение", "conclusion",
    "вывод", "summary", "итог", "implications", "what this means",
]


def _is_heading(para) -> tuple[bool, int]:
    """Return (is_heading, level) from paragraph style."""
    style_name = (para.style.name or "").lower()
    if style_name.startswith("heading 1"):
        return True, 1
    if style_name.startswith("heading 2"):
        return True, 2
    if style_name.startswith("heading 3"):
        return True, 3
    return False, 0


def _normalize_heading(text: str) -> str:
    return text.strip().strip("*#").strip()


# ── main loader ──────────────────────────────────────────────────────────────

def load(path: Path) -> LLMReportIR:
    doc = Document(str(path))
    all_text = "\n".join(_para_text(p) for p in doc.paragraphs)

    detected_agent = _detect_agent(all_text)
    cite_format = _detect_cite_format(all_text)
    cite_regex = _get_cite_regex(cite_format)

    sections: list[RawSection] = []
    citations_raw: dict[int, RawCitation] = {}
    open_questions: list[str] = []
    parser_notes: list[str] = []

    current_section: RawSection | None = None
    in_sources = False
    in_open_questions = False

    for para in doc.paragraphs:
        text = _para_text(para).strip()
        if not text:
            continue

        # Inline ``[N]`` hyperlink markers can appear anywhere in the body
        # (ChatGPT Deep Research style — no trailing Sources block).
        _extract_inline_citations(para, citations_raw)

        is_hd, level = _is_heading(para)

        if is_hd:
            heading = _normalize_heading(text)
            heading_low = heading.lower()

            # Detect special sections
            if any(h in heading_low for h in _SOURCES_HINTS):
                in_sources = True
                in_open_questions = False
                current_section = None
                continue

            if any(h in heading_low for h in _OPEN_QUESTIONS_HINTS):
                in_open_questions = True
                in_sources = False
                current_section = None
                continue

            if any(h in heading_low for h in _SKIP_SECTIONS):
                in_sources = False
                in_open_questions = False
                current_section = None
                parser_notes.append(f"Skipped section '{heading}' (meta/conclusions)")
                continue

            # Regular content section
            in_sources = False
            in_open_questions = False
            current_section = RawSection(heading=heading, level=level)
            sections.append(current_section)
            continue

        # Sources block — parse citation lines
        if in_sources:
            _parse_citation_para(para, text, cite_format, citations_raw)
            continue

        # Open questions block — collect bullet points
        if in_open_questions:
            clean = re.sub(r"^[-•*]\s*", "", text).strip()
            if clean:
                open_questions.append(clean)
            continue

        # Skip if we're in a skipped section (current_section is None after skip)
        if current_section is None:
            continue

        # Regular paragraph — add to current section, parse inline cite markers
        current_section.paragraphs.append(text)

    # Also parse tables
    _parse_tables(doc, sections)

    # Build ordered citation list
    citations = [citations_raw[k] for k in sorted(citations_raw)]

    # Fallback: if no structured citations found, scan all text for URLs
    if not citations:
        parser_notes.append("No structured citation block found; scanning full text for URLs")
        citations = _extract_citations_from_text(all_text, cite_regex)

    return LLMReportIR(
        source_filename=path.name,
        detected_agent=detected_agent,
        detected_cite_format=cite_format,
        sections=sections,
        citations=citations,
        open_questions=open_questions,
        parser_notes=parser_notes,
    )


def _get_cite_regex(fmt: str) -> re.Pattern:
    if fmt == "bracket_n":
        return _RE_BRACKET_N
    if fmt == "paren_source_n":
        return _RE_PAREN_N
    return _RE_BRACKET_N  # default fallback


def _parse_citation_para(
    para,
    text: str,
    fmt: str,
    out: dict[int, RawCitation],
) -> None:
    """Try to parse a single paragraph as a citation entry."""
    # Try to match "[N] Title — Publisher — URL" style
    m = _RE_CITE_LINE_BRACKET.match(text)
    if m:
        cite_id = int(m.group(1))
        label = (m.group(2) or "").strip()
        url = m.group(3) or _get_para_url(para) or ""
        if not url:
            url_m = _RE_URL.search(text)
            url = url_m.group(0) if url_m else ""
        out[cite_id] = RawCitation(
            cite_id=cite_id,
            raw_marker=f"[{cite_id}]",
            url=url.rstrip(".,)>"),
            title=label[:200],
        )
        return

    # Numbered list style without brackets: "1. Title URL"
    m2 = re.match(r"^(\d+)[.)]\s+(.+)$", text)
    if m2:
        cite_id = int(m2.group(1))
        rest = m2.group(2)
        url_m = _RE_URL.search(rest)
        url = url_m.group(0).rstrip(".,)>") if url_m else _get_para_url(para)
        title = rest[: url_m.start()].strip() if url_m else rest[:200]
        out[cite_id] = RawCitation(
            cite_id=cite_id,
            raw_marker=str(cite_id),
            url=url,
            title=title[:200],
        )
        return

    # Hyperlink-only paragraph inside sources block
    url = _get_para_url(para)
    if url and out:
        # Associate with next sequential id
        next_id = max(out.keys()) + 1
        out[next_id] = RawCitation(
            cite_id=next_id,
            raw_marker=str(next_id),
            url=url,
            title=text[:200],
        )


def _extract_citations_from_text(
    text: str, regex: re.Pattern
) -> list[RawCitation]:
    """Fallback: collect all unique [N] markers and URLs from raw text."""
    ids_seen: set[int] = set()
    result: list[RawCitation] = []
    for m in regex.finditer(text):
        cid = int(m.group(1))
        if cid not in ids_seen:
            ids_seen.add(cid)
            result.append(RawCitation(cite_id=cid, raw_marker=m.group(0), url=""))
    return result


def _parse_tables(doc, sections: list[RawSection]) -> None:
    """Append table rows to the last content section before the table in the doc."""
    if not sections:
        return
    for table in doc.tables:
        target = sections[-1]  # best-effort: attach to last section
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                target.table_rows.append(cells)
